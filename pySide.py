from PySide import QtGui
from PySide import QtCore
import sys
import requests
import calendar
import time
from docx import Document
from openpyxl import load_workbook
from datetime import datetime as dt

#Unique API ID used to download forecast from web service
API_URL = "http://api.openweathermap.org/data/2.5/forecast?q="
API_ID = "&APPID=4c8807c1a33b7cbd1f3f04a2f03a0bf3"

#Class with http codes to identify web error
class HttpCodes:
    ok = 200
    bad_request = 400
    unauthorized = 401
    not_found = 404

#Abstract class to provide service to download data from web server or excel file
class HistoricalWeatherProvider:

    def get_historical_weather_for_city(self, city_name):
        raise NotImplementedError

#Weather client to download temperature from web service
class OWMForecastWeatherClient(HistoricalWeatherProvider):

    def __init__(self, city_name, end):
        self.city = city_name
        self.end_date = end

    #Get weather from web service for five days maximum or if user check earlier date on calendar
    #it should be used as end date
    def get_historical_weather_for_city(self):

        data_to_export = {}
        #Provided by user date convert to datetime
        user_end_date = calendar.timegm(time.strptime(self.end_date, '%d.%m.%Y'))

        #For every city provided by user download data for them from web service
        for every_city in self.city:
            url = API_URL + every_city +"&units=metric" + API_ID
            response = requests.get(url)

            #If request is correct parse json
            if HttpCodes.ok == response.status_code:
                json_string = response.json()
                list_length = len(json_string['list'])
                weather_list = []
                temp_max_list = []
                temp_min_list = []
                for every_date in range(list_length):
                    if(user_end_date >= json_string['list'][every_date]['dt']):
                        weather_list.append(json_string['list'][every_date]['main']['temp'])
                        temp_max_list.append(json_string['list'][every_date]['main']['temp_max'])
                        temp_min_list.append(json_string['list'][every_date]['main']['temp_min'])
                        end_date = json_string['list'][every_date]['dt_txt']
                        start_date = json_string['list'][0]['dt_txt']
                print("Sucessfully downloaded temperature from service")
            else:
                print(response.status_code)
                return None

            #Return data as dictionary of city and weather details - it will be used to generate word report
            data_to_export.update({every_city: {'Average temperature': sum(weather_list)/list_length,
                                                'Temp_max': max(temp_max_list),
                                                'Temp_min': min(temp_min_list),
                                                'End date': end_date,
                                                'Start date': start_date}})
        print("Data to export: \n", data_to_export)
        return data_to_export

#Weather client to load excel file and parse data
class XlsxHistoricalWeatherReader(HistoricalWeatherProvider):

    def __init__(self, start, end):
        self.start_dat = start
        self.end_date = end
        self.city_list = []
        self.date_list = []
        self.temperature = []

    #Get weather from excel file for the days user check on calendar
    #fileLoaction - source of file path
    def get_historical_weather_for_city(self, file_location):

        #Load excel file and make it active
        wb = load_workbook(file_location[0])
        ws = wb.active

        #start and end date choosen on calendar converted to datetime
        start_date = dt.strptime(self.start_dat, "%d.%m.%Y")
        end_date = dt.strptime(self.end_date, "%d.%m.%Y")

        #Iteration on rows in excel sheet to get date
        for row in ws.iter_rows(min_row=2, max_col=1, max_row=10):
            for cell in row:
                current_date = dt.strptime(cell.value, "%d.%m.%Y")
                #If cell is empty don't add them on the list
                if (cell.value is not None):
                    #If current date in cell is it not in range of start and end date provided by user
                    #don't add them on the list
                    if current_date >= start_date and current_date <= end_date:
                        self.date_list.append(cell.value)
                else:
                    break
        if self.date_list:
            #Iteration in excell sheet on cells to get city name
            for row in ws.iter_rows(max_row=1, min_col=2, max_col=10):
                for cell in row:
                    if (cell.value is not None):
                        self.city_list.append(cell.value)
                    else:
                        break

            #Start column should be 2, because there is first city in row
            start_column = 2
            weather_list = []
            data_to_export = {}

            #For every city iter by row and cells to get weather details
            for city in self.city_list:
                for row in ws.iter_rows(min_row=2, min_col=start_column, max_col = start_column, max_row= len(self.date_list)):
                    for cell in row:
                        if (cell.value is not None):
                            weather_list.append(cell.value)
                        else:
                            break
                start_column += 1
                #Return data as dictionary of city and weather details - it will be used to generate word report
                data_to_export.update({city: {'Average temperature': sum(weather_list)/len(weather_list),
                                                'Temp_max': max(weather_list),
                                                'Temp_min': min(weather_list),
                                                'End date': self.date_list[-1],
                                                'Start date': self.date_list[0]}})
                weather_list.clear()
            print("Data to export from excel: \n", data_to_export)
            return data_to_export
        else:
            return False

#Docx client to generate word report from excel file or web service
class DocxHistoricalWeatherReporter:

    def __init__(self, report_source):
        self.report_path = report_source

    #Generate report based on weather_data - dictionary of city and weather details
    def generateReport(self, weather_data):

        #Create word document with header
        document = Document()
        document.add_heading('Temperature Report', 0)

        #If weather_data is not empty generate word
        if weather_data is not None:
            for every_city in weather_data:
                document.add_heading(every_city, level=1)
                table = document.add_table(rows=2, cols=5)

                #Get data from dictionary for every city
                temperature_of_city = weather_data[every_city]['Average temperature']
                end_date = weather_data[every_city]['End date']
                start_date = weather_data[every_city]['Start date']
                temp_max = weather_data[every_city]['Temp_max']
                temp_min = weather_data[every_city]['Temp_min']

                #Create header in word table
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'From'
                hdr_cells[1].text = 'To'
                hdr_cells[2].text = 'Average'
                hdr_cells[3].text = 'Max'
                hdr_cells[4].text = 'Min'

                #Put data in word ttable
                hdr_data = table.rows[1].cells
                hdr_data[0].text = str(start_date)
                hdr_data[1].text = str(end_date)
                hdr_data[2].text = str(round(temperature_of_city))
                hdr_data[3].text = str(temp_max)
                hdr_data[4].text = str(temp_min)
        else:
            print("Couldn't generate word report")
            return None

        document.add_page_break()
        document.save(self.report_path)

        print("Word raport successfully generated from Web Service and Excel")

#Generate GUI for weather provider
class GuiWeatherProvider(QtGui.QWidget):

    def __init__(self):
        super(GuiWeatherProvider, self).__init__()

        self.initUI()
        self.city_list_to_export = []
        self.file_location = ''
        self.date_string_to_web_service = ''
        self.start_date_string_to_excel = ''
        self.end_date_string_to_excel = ''

    def initUI(self):

        self.st = QtGui.QLabel("Generate Temperature Report From Web Service \n\nCity:")

        self.text_box = QtGui.QTextEdit(self)
        self.text_box.setMaximumSize(150, 30)

        self.button_add_city = QtGui.QPushButton('Add City', self)

        self.city_list = QtGui.QListWidget(self)
        self.city_list.readOnly = True

        self.button_get_from_web = QtGui.QPushButton('Get Temperature From Web', self)

        self.calendar_for_web_service = QtGui.QCalendarWidget()

        self.st2 = QtGui.QLabel("\n Generate Temperature Report From Excel\n")

        self.button_browse_file = QtGui.QPushButton('Browse File', self)

        self.text_box_file_location = QtGui.QTextEdit(self)
        self.text_box_file_location.setMaximumSize(550, 40)

        self.star_date_label = QtGui.QLabel("Start Date")
        self.star_date_label.setAlignment(QtCore.Qt.AlignCenter)

        self.end_date_label = QtGui.QLabel("End Date")
        self.end_date_label.setAlignment(QtCore.Qt.AlignCenter)

        self.dialog = QtGui.QFileDialog(self)

        self.button_get_from_excel = QtGui.QPushButton('Get Temperature From Excel', self)

        self.calendar_for_excel_start_date = QtGui.QCalendarWidget()

        self.calendar_for_excel_end_date = QtGui.QCalendarWidget()

        self.resize(500, 600)
        self.center()
        self.setWindowTitle('Center')

        layout = QtGui.QGridLayout()
        layout.addWidget(self.st, 0, 0)
        layout.addWidget(self.text_box, 1, 0)
        layout.addWidget(self.button_add_city, 2, 0)
        layout.addWidget(self.city_list, 3, 0)
        layout.addWidget(self.calendar_for_web_service, 4, 0)
        layout.addWidget(self.button_get_from_web, 5, 0)
        layout.addWidget(self.st2, 6, 0)
        layout.addWidget(self.button_browse_file, 7, 0)
        layout.addWidget(self.text_box_file_location, 8, 0)
        layout.addWidget(self.star_date_label, 9, 0)
        layout.addWidget(self.end_date_label, 9, 1)
        layout.addWidget(self.calendar_for_excel_start_date, 10, 0)
        layout.addWidget(self.calendar_for_excel_end_date, 10, 1)
        layout.addWidget(self.button_get_from_excel)

        self.setLayout(layout)

        #Connect buttons to button state - if button is clicked do defined function
        self.button_add_city.clicked.connect(self.btnstateForCityButton)
        self.button_get_from_web.clicked.connect(self.btnstateForWebService)
        self.button_browse_file.clicked.connect(self.btnstateBrowseButton)
        self.button_get_from_excel.clicked.connect(self.btnstateForExcel)

        #Connect calendar date picker with changes (On click)
        self.connect(self.calendar_for_web_service, QtCore.SIGNAL('selectionChanged()'), self.date_changed_web_service)
        self.connect(self.calendar_for_excel_start_date, QtCore.SIGNAL('selectionChanged()'), self.date_changed_excel)
        self.connect(self.calendar_for_excel_end_date, QtCore.SIGNAL('selectionChanged()'), self.date_changed_excel)

    #If date on calendar for web service is changed set new date
    def date_changed_web_service(self):
        date = self.calendar_for_web_service.selectedDate()
        self.date_string_to_web_service = str(date.day()) + '.' + str(date.month()) + '.' + str(date.year())
        print(self.date_string_to_web_service)

    #If date on calendar for excel data is changed set new date
    def date_changed_excel(self):
        start_date = self.calendar_for_excel_start_date.selectedDate()
        end_date = self.calendar_for_excel_end_date.selectedDate()

        self.start_date_string_to_excel = str(start_date.day()) + '.' + str(start_date.month()) + '.' + str(start_date.year())
        self.end_date_string_to_excel = str(end_date.day()) + '.' + str(end_date.month()) + '.' + str(end_date.year())

    #If button to add new city is clicked add city on the list or if text box is empty display message
    def btnstateForCityButton(self):
        if self.button_add_city.isChecked() == False:
            city = self.text_box.toPlainText()
            if city:
                city = self.text_box.toPlainText()
                self.city_list.addItem(city)
                self.city_list_to_export.append(city)
                self.text_box.clear()
            else:
                msgBox = QtGui.QMessageBox()
                msgBox.setText('Please input City')
                msgBox.show()
                msgBox.exec_()

    #If button to download data from web service is clicked download data from server and generate word report
    def btnstateForWebService(self):
        if self.button_get_from_web.isChecked() == False:
            if self.city_list_to_export:
                if self.date_string_to_web_service:
                    #Download data from server
                    weather_from_internet = OWMForecastWeatherClient(self.city_list_to_export, self.date_string_to_web_service)
                    temperature_of_cities = weather_from_internet.get_historical_weather_for_city()
                    #Generate word report
                    docx = DocxHistoricalWeatherReporter("C:/Users/krzacjan/Desktop/pythonHomeWork/temperature_report_from_web.docx")
                    docx.generateReport(temperature_of_cities)

                    msgBox = QtGui.QMessageBox()
                    msgBox.setText('Raport From Web Service Was Generated')
                    msgBox.show()
                    msgBox.exec_()
                else:
                    msgBox = QtGui.QMessageBox()
                    msgBox.setText('Please choose end day from calendar')
                    msgBox.show()
                    msgBox.exec_()
            else:
                msgBox = QtGui.QMessageBox()
                msgBox.setText('Please add City to list')
                msgBox.show()
                msgBox.exec_()

    #If button to download data from excel file is clicked download data from excel sheet and generate word report
    def btnstateForExcel(self):
        if self.button_get_from_excel.isChecked() == False:
            if self.file_location:
                if self.start_date_string_to_excel and self.end_date_string_to_excel:
                    #Download data from excel sheet
                    weather_from_excel = XlsxHistoricalWeatherReader(self.start_date_string_to_excel, self.end_date_string_to_excel)
                    temperature_of_cities_from_excel = weather_from_excel.get_historical_weather_for_city(self.file_location)
                    #Generate word report
                    if temperature_of_cities_from_excel:
                        docx = DocxHistoricalWeatherReporter("C:/Users/krzacjan/Desktop/pythonHomeWork/temperature_report_from_excel.docx")
                        docx.generateReport(temperature_of_cities_from_excel)

                        msgBox = QtGui.QMessageBox()
                        msgBox.setText('Raport From Excel Was Generated')
                        msgBox.show()
                        msgBox.exec_()
                    else:
                        msgBox = QtGui.QMessageBox()
                        msgBox.setText('Please choose date in range')
                        msgBox.show()
                        msgBox.exec_()
                else:
                    msgBox = QtGui.QMessageBox()
                    msgBox.setText('Please choose end day from calendar')
                    msgBox.show()
                    msgBox.exec_()
            else:
                msgBox = QtGui.QMessageBox()
                msgBox.setText('Please input File Location')
                msgBox.show()
                msgBox.exec_()

    #If button to browse file path is clicked open dialog to find file path
    def btnstateBrowseButton(self):
        if self.button_browse_file.isChecked() == False:
            self.file_location = self.dialog.getOpenFileName()
            self.text_box_file_location.setText(str(self.file_location[0]))

    def center(self):
        qr = self.frameGeometry()
        cp = QtGui.QDesktopWidget().availableGeometry().center()
        qr.moveCenter(cp)
        self.move(qr.topLeft())

def main():
    app = QtGui.QApplication(sys.argv)
    ex = GuiWeatherProvider()
    ex.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()