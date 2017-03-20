
import requests
import calendar
import time
import datetime
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl import load_workbook

API_URL = "http://api.openweathermap.org/data/2.5/forecast?q="
API_ID = "&APPID=4c8807c1a33b7cbd1f3f04a2f03a0bf3"

class HttpCodes:
    ok = 200
    bad_request = 400
    unauthorized = 401
    not_found = 404

class HistoricalWeatherProvider:

    def getHistoricalWeatherForCity(self, city_name):
        raise NotImplementedError

    def getAllHistoricalTemperatures(self):
        raise NotImplementedError

class OWMHistoricalWeatherClient(HistoricalWeatherProvider):

    def __init__(self, city_name, end):
        self.city = city_name
        self.end_date = end

    def getHistoricalWeatherForFiveDaysForCity(self):

        data_to_export = {}

        user_end_date = calendar.timegm(time.strptime(self.end_date, '%d.%m.%Y'))

        for every_city in self.city:
            url = API_URL + every_city +"&units=metric" + API_ID
            print(url)
            response = requests.get(url)

            if HttpCodes.ok == response.status_code:
                json_string = response.json()
                list_length = len(json_string['list'])
                weather_list = []
                temp_max_list = []
                temp_min_list = []
                for every_date in range(list_length):

                    if(user_end_date >= json_string['list'][every_date]['dt']):
                        weather_list.append(json_string['list'][every_date]['main']['temp'])
                        temp_max_list.append(json_string['list'][every_date]['main']['temp_max'])
                        temp_min_list.append(json_string['list'][every_date]['main']['temp_min'])
                        end_date = json_string['list'][every_date]['dt_txt']
                        start_date = json_string['list'][0]['dt_txt']
                print("Sucessfully downloaded temperature from service")

            else:
                print(response.status_code)
                return None

            data_to_export.update({every_city: {'Average temperature': sum(weather_list)/list_length,
                                                'Temp_max': max(temp_max_list),
                                                'Temp_min': min(temp_min_list),
                                                'End date': end_date,
                                                'Start date': start_date}})
        print("Data to export: \n", data_to_export)
        return data_to_export

class XlsxHistoricalWeatherReader(HistoricalWeatherProvider):

    def __init__(self, city_name, start, end):
        self.city = city_name
        self.start_dat = start
        self.end_date = end
        self.city_list = []
        self.date_list = []
        self.temperature = []

    def getHistoricalWeatherForCity(self):
        wb = load_workbook('temperature.xlsx')
        ws = wb.active

        for row in ws.iter_rows(min_row=2, max_col=1, max_row=10):
            for cell in row:
                if (cell.value is not None):
                    self.date_list.append(cell.value)
                else:
                    break

        for row in ws.iter_rows(max_row=1, min_col=2, max_col=10):
            for cell in row:
                if (cell.value is not None):
                    self.city_list.append(cell.value)
                else:
                    break

        start_column = 2
        weather_list = []
        data_to_export = {}
        for city in self.city_list:
            for row in ws.iter_rows(min_row=2, min_col=start_column, max_col = start_column, max_row=10):
                for cell in row:
                    if (cell.value is not None):
                        weather_list.append(cell.value)
                    else:
                        break
            start_column += 1
            data_to_export.update({city: {'Average temperature': sum(weather_list)/len(weather_list),
                                            'Temp_max': max(weather_list),
                                            'Temp_min': min(weather_list),
                                            'End date': self.date_list[-1],
                                            'Start date': self.date_list[0]}})
            weather_list.clear()
        print("Data to export from excel: \n", data_to_export)
        return data_to_export


class DocxHistoricalWeatherReporter:

    def __init__(self, report_source):
        self.report_path = report_source

    def generateReportFromWeb(self, weather_data):
        document = Document()

        document.add_heading('Temperature Report', 0)

        if weather_data is not None:
            for every_city in weather_data:
                document.add_heading(every_city, level=1)
                table = document.add_table(rows=2, cols=5)

                temperature_of_city = weather_data[every_city]['Average temperature']
                end_date = weather_data[every_city]['End date']
                start_date = weather_data[every_city]['Start date']
                temp_max = weather_data[every_city]['Temp_max']
                temp_min = weather_data[every_city]['Temp_min']

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'From'
                hdr_cells[1].text = 'To'
                hdr_cells[2].text = 'Average'
                hdr_cells[3].text = 'Max'
                hdr_cells[4].text = 'Min'

                hdr_data = table.rows[1].cells
                hdr_data[0].text = str(start_date)
                hdr_data[1].text = str(end_date)
                hdr_data[2].text = str(round(temperature_of_city))
                hdr_data[3].text = str(temp_max)
                hdr_data[4].text = str(temp_min)
        else:
            print("Couldn't download temperature from service")
            return None

        document.add_page_break()

        document.save(self.report_path)

        print("Word raport successfully generated from Web Service and Excel")

def main():

    city_list = []

    weather_from_internet = OWMHistoricalWeatherClient({"Wroclaw", "Berlin"}, "23.03.2017")
    temperature_of_cities = weather_from_internet.getHistoricalWeatherForFiveDaysForCity()

    docx = DocxHistoricalWeatherReporter("C:/Users/krzacjan/Desktop/pythonHomeWork/temperature_report_from_web.docx")
    docx.generateReportFromWeb(temperature_of_cities)

    weather_from_excel =  XlsxHistoricalWeatherReader({"Wroclaw", "Berlin"}, "01.12.2016", "05.12.2016")
    temperature_of_cities_from_excel = weather_from_excel.getHistoricalWeatherForCity()

    docx = DocxHistoricalWeatherReporter("C:/Users/krzacjan/Desktop/pythonHomeWork/temperature_report_from_excel.docx")
    docx.generateReportFromWeb(temperature_of_cities_from_excel)

if __name__ == "__main__":
    main()



